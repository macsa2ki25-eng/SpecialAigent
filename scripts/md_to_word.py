#!/usr/bin/env python3
"""
SpecialAigent - Markdown → Word変換スクリプト
テスト問題・指導案・校務文書をWord形式で出力する

使い方:
    python3 scripts/md_to_word.py input.md output.docx
    python3 scripts/md_to_word.py input.md               # → input.docx として出力
"""

import sys
import re
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT


def create_styled_document():
    """学校文書向けのスタイル設定済みDocumentを作成"""
    doc = Document()

    # ページ設定（A4、余白2cm）
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # デフォルトフォント
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Yu Mincho'  # 游明朝（Windows標準）
    font.size = Pt(10.5)

    return doc


def parse_markdown_to_word(md_text, doc):
    """Markdownテキストを解析してWord文書に変換する"""
    lines = md_text.split('\n')
    i = 0
    in_table = False
    table_rows = []
    in_code_block = False
    code_lines = []

    while i < len(lines):
        line = lines[i]

        # コードブロック
        if line.strip().startswith('```'):
            if in_code_block:
                # コードブロック終了 → テキストボックス風に出力
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph()
                p.style = doc.styles['Normal']
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                p.paragraph_format.left_indent = Cm(1.0)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # テーブル処理
        if '|' in line and line.strip().startswith('|'):
            cells = [c.strip() for c in line.strip().strip('|').split('|')]
            # セパレータ行（---）はスキップ
            if all(re.match(r'^:?-+:?$', c) for c in cells):
                i += 1
                continue
            if not in_table:
                in_table = True
                table_rows = []
            table_rows.append(cells)
            i += 1
            continue
        elif in_table:
            # テーブル終了 → Word表に変換
            _add_table(doc, table_rows)
            in_table = False
            table_rows = []
            # この行はテーブルではないので、fallthrough して処理を続ける

        # 空行
        if not line.strip():
            i += 1
            continue

        # 水平線
        if line.strip() in ('---', '***', '___'):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            # 薄い線として罫線を追加
            run = p.add_run('─' * 60)
            run.font.size = Pt(6)
            run.font.color.rgb = RGBColor(200, 200, 200)
            i += 1
            continue

        # 見出し
        if line.startswith('# '):
            heading = doc.add_heading(clean_md(line[2:]), level=0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in heading.runs:
                run.font.size = Pt(16)
            i += 1
            continue
        if line.startswith('## '):
            heading = doc.add_heading(clean_md(line[3:]), level=1)
            for run in heading.runs:
                run.font.size = Pt(14)
            i += 1
            continue
        if line.startswith('### '):
            heading = doc.add_heading(clean_md(line[4:]), level=2)
            for run in heading.runs:
                run.font.size = Pt(12)
            i += 1
            continue

        # 通常の段落
        p = doc.add_paragraph()
        _add_formatted_text(p, line)
        i += 1

    # 最後にテーブルが残っていた場合
    if in_table and table_rows:
        _add_table(doc, table_rows)


def _add_table(doc, rows):
    """テーブルデータをWord表に変換"""
    if not rows:
        return

    num_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for ri, row in enumerate(rows):
        for ci, cell_text in enumerate(row):
            if ci < num_cols:
                cell = table.cell(ri, ci)
                cell.text = ''
                p = cell.paragraphs[0]
                _add_formatted_text(p, cell_text)
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)

                # ヘッダー行（最初の行）は太字
                if ri == 0:
                    for run in p.runs:
                        run.bold = True

    doc.add_paragraph()  # テーブル後にスペース


def _add_formatted_text(paragraph, text):
    """Markdownのインライン書式をWordに変換"""
    text = text.strip()

    # パターン: **太字**, *斜体*, `コード`
    pattern = r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`|([^*`]+))'
    matches = re.finditer(pattern, text)

    for match in matches:
        if match.group(2):  # **太字**
            run = paragraph.add_run(match.group(2))
            run.bold = True
        elif match.group(3):  # *斜体*
            run = paragraph.add_run(match.group(3))
            run.italic = True
        elif match.group(4):  # `コード`
            run = paragraph.add_run(match.group(4))
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
        elif match.group(5):  # 通常テキスト
            paragraph.add_run(match.group(5))


def clean_md(text):
    """Markdownの装飾を除去してプレーンテキストにする"""
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    return text.strip()


def convert(input_path, output_path=None):
    """メイン変換処理"""
    if output_path is None:
        output_path = os.path.splitext(input_path)[0] + '.docx'

    with open(input_path, 'r', encoding='utf-8') as f:
        md_text = f.read()

    doc = create_styled_document()
    parse_markdown_to_word(md_text, doc)
    doc.save(output_path)

    return output_path


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("使い方: python3 md_to_word.py <入力.md> [出力.docx]")
        sys.exit(1)

    input_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else None
    result = convert(input_file, output_file)
    print(f"Word文書を生成しました: {result}")
